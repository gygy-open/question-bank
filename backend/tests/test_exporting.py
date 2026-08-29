"""导出子系统单测:LaTeX 转义/遍历、答案行内化、装配(分节/编号/位置)。"""

import types

from app.schemas.paper import ContentPosition, OutputFormat
from app.services.exporting.answer import answer_spec_to_inline
from app.services.exporting.assemble import PaperAssembler
from app.services.exporting.contracts import ExportOption, ExportOptions
from app.services.exporting.richdoc.latex import (
    latex_escape,
    rich_doc_to_latex,
    rich_inline_to_latex,
)
from app.services.question_content import to_db_json
from app.services.question_content_converter import markdown_to_rich_doc


# --------------------------------------------------------------------------- #
# LaTeX 转义与遍历
# --------------------------------------------------------------------------- #
def test_latex_escape_specials():
    assert latex_escape("a & b % c $ d # e _ f { g } h ~ i ^ j") == (
        r"a \& b \% c \$ d \# e \_ f \{ g \} h \textasciitilde{} i \textasciicircum{} j"
    )


def test_latex_escape_backslash_first():
    # 反斜杠先转义,不能污染后续注入的命令。
    assert latex_escape(r"\%") == r"\textbackslash{}\%"


def test_rich_doc_to_latex_marks_and_math():
    doc = markdown_to_rich_doc("**粗** 和 *斜* 与 $x^2$")
    tex = rich_doc_to_latex(doc)
    assert r"\textbf{粗}" in tex
    assert r"\textit{斜}" in tex
    assert "$x^2$" in tex
    assert '"type"' not in tex


def test_rich_doc_to_latex_text_is_escaped():
    doc = markdown_to_rich_doc("100% 正确 & 完整")
    tex = rich_doc_to_latex(doc)
    assert r"\%" in tex
    assert r"\&" in tex


def test_rich_doc_to_latex_unknown_node_keeps_text():
    doc = {"type": "doc", "content": [{"type": "weirdBlock", "text": "留住"}]}
    assert "留住" in rich_doc_to_latex(doc)


def test_rich_doc_to_latex_empty():
    assert rich_doc_to_latex(None) == ""
    assert rich_doc_to_latex({"type": "doc", "content": []}) == ""


def test_rich_doc_to_latex_blockquote():
    doc = {
        "type": "doc",
        "content": [{"type": "blockquote", "content": [
            {"type": "paragraph", "content": [{"type": "text", "text": "引文"}]}
        ]}],
    }
    tex = rich_doc_to_latex(doc)
    assert r"\begin{quote}" in tex and r"\end{quote}" in tex
    assert "引文" in tex


def test_rich_doc_to_latex_code_block_verbatim():
    doc = {
        "type": "doc",
        "content": [{"type": "codeBlock", "content": [{"type": "text", "text": "a & b % c"}]}],
    }
    tex = rich_doc_to_latex(doc)
    assert r"\begin{verbatim}" in tex and r"\end{verbatim}" in tex
    # verbatim 原样保留，不转义特殊字符。
    assert "a & b % c" in tex


def test_rich_doc_to_latex_horizontal_rule():
    doc = {"type": "doc", "content": [{"type": "horizontalRule"}]}
    assert r"\rule{" in rich_doc_to_latex(doc)


def test_docx_renderer_blockquote_codeblock_hr():
    from docx import Document

    from app.services.exporting.richdoc.docx import DocxRichRenderer

    document = Document()
    DocxRichRenderer().render_doc(
        document,
        {
            "type": "doc",
            "content": [
                {"type": "blockquote", "content": [
                    {"type": "paragraph", "content": [{"type": "text", "text": "引文"}]}
                ]},
                {"type": "codeBlock", "content": [{"type": "text", "text": "print(1)"}]},
                {"type": "horizontalRule"},
            ],
        },
    )
    texts = "\n".join(p.text for p in document.paragraphs)
    assert "引文" in texts
    assert "print(1)" in texts
    # 分隔线以末段底边框呈现。
    assert "w:pBdr" in document.paragraphs[-1]._p.xml


# --------------------------------------------------------------------------- #
# 答案行内化
# --------------------------------------------------------------------------- #
def test_answer_inline_single_choice_uses_label():
    opts = [ExportOption(id="opt_a", label="A", content=None), ExportOption(id="opt_b", label="B", content=None)]
    nodes = answer_spec_to_inline({"kind": "single_choice", "correct": "opt_b"}, opts)
    assert rich_inline_to_latex(nodes) == "B"


def test_answer_inline_true_false():
    nodes = answer_spec_to_inline({"kind": "true_false", "correct": False}, [])
    assert rich_inline_to_latex(nodes) == "错"


def test_answer_inline_fill_preserves_math():
    answer = {
        "kind": "fill_in_the_blank",
        "blanks": [{"id": "blk_1", "accept": [markdown_to_rich_doc("$x^2$"), markdown_to_rich_doc("4")]}],
    }
    tex = rich_inline_to_latex(answer_spec_to_inline(answer, []))
    assert "$x^2$" in tex
    assert "或" in tex
    assert "4" in tex


# --------------------------------------------------------------------------- #
# 装配:分节 / 编号 / 位置
# --------------------------------------------------------------------------- #
def _q(qid, content="题", qtype="single_choice"):
    return types.SimpleNamespace(
        id=qid,
        q_type=qtype,
        difficulty=1,
        content=to_db_json(markdown_to_rich_doc(content)),
        options=None,
        answer=to_db_json({"kind": "true_false", "correct": True}) if qtype == "true_false" else None,
        thinking=to_db_json(markdown_to_rich_doc("思路")),
        analysis=to_db_json(markdown_to_rich_doc("解析")),
        summary=None,
        source="来源.docx",
    )


def _opts(**kw):
    base = dict(title="卷子")
    base.update(kw)
    return ExportOptions(**base)


def test_assemble_continuous_numbering_across_sections():
    qs = [_q(1), _q(2), _q(3)]
    doc = PaperAssembler(_opts()).assemble(qs, ["第一部分", None, "第二部分"])
    numbers = [q.number for s in doc.sections for q in s.questions]
    assert numbers == [1, 2, 3]
    assert [s.title for s in doc.sections] == ["第一部分", "第二部分"]


def test_assemble_after_question_shows_inline_details():
    doc = PaperAssembler(_opts()).assemble([_q(1, qtype="true_false")], [None])
    q = doc.sections[0].questions[0]
    assert q.answer is not None
    assert q.thinking is not None
    assert doc.has_appendix is False


def test_assemble_end_of_paper_moves_details_to_appendix():
    opts = _opts()
    opts.details_at_end = True
    doc = PaperAssembler(opts).assemble([_q(1, qtype="true_false")], [None])
    body_q = doc.sections[0].questions[0]
    assert body_q.answer is None and body_q.thinking is None
    assert doc.has_appendix is True
    appendix_q = doc.appendix[0].questions[0]
    assert appendix_q.number == 1
    assert appendix_q.answer is not None


def test_assemble_hidden_drops_all_details():
    opts = _opts()
    opts.hidden_details = True
    doc = PaperAssembler(opts).assemble([_q(1, qtype="true_false")], [None])
    q = doc.sections[0].questions[0]
    assert q.answer is None and q.thinking is None and q.analysis is None
    assert doc.has_appendix is False


def test_assemble_include_flags_filter_fields():
    opts = _opts(include_analysis=False, include_source=True)
    doc = PaperAssembler(opts).assemble([_q(1, qtype="true_false")], [None])
    q = doc.sections[0].questions[0]
    assert q.thinking is None      # include_analysis=False
    assert q.analysis is not None  # include_explanation 默认 True
    assert q.source == "来源.docx"


def test_assemble_free_response_reserves_space():
    doc = PaperAssembler(_opts()).assemble([_q(1, qtype="free_response")], [None])
    assert doc.sections[0].questions[0].reserve_space is True

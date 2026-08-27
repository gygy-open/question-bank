"""组稿导出渲染器冒烟测试:CompositionExportDoc → 磁盘文件(docx/latex zip),不校验具体排版。"""

import os
import zipfile

from docx import Document

from app.schemas.paper import OutputFormat
from app.services.exporting.composition_assemble import CompositionAssembler
from app.services.exporting.composition_contracts import (
    CompositionExportDoc,
    ExportAnswerEntry,
    ExportHeadingNode,
    ExportOption,
    ExportPageBreakNode,
    ExportQuestionDetailsNode,
    ExportQuestionNode,
    ExportRichTextNode,
)
from app.services.exporting.composition_registry import composition_renderer_for
from app.services.exporting.renderers.composition_docx import CompositionDocxRenderer
from app.services.exporting.renderers.composition_latex import CompositionLatexRenderer


def _rich_doc(text: str) -> dict:
    return {"type": "doc", "content": [{"type": "paragraph", "content": [{"type": "text", "text": text}]}]}


def _rich_doc_aligned(text: str, align: str) -> dict:
    return {
        "type": "doc",
        "content": [{"type": "paragraph", "attrs": {"textAlign": align}, "content": [{"type": "text", "text": text}]}],
    }


def _sample_doc() -> CompositionExportDoc:
    options = [
        ExportOption(id="opt_a", label="A", content=_rich_doc("2")),
        ExportOption(id="opt_b", label="B", content=_rich_doc("3")),
    ]
    question = ExportQuestionNode(
        number="1",
        score=4.5,
        q_type="single_choice",
        stem=_rich_doc("1+1=?"),
        options=options,
        option_columns=2,
        answer={"kind": "single_choice", "correct": "opt_a"},
        thinking=_rich_doc("先算加法"),
        analysis=_rich_doc("因为 1+1=2"),
        summary=None,
    )
    module = ExportQuestionDetailsNode(
        scope="all",
        children=[
            ExportHeadingNode(level=3, content=_rich_doc("答案汇总")),
            ExportAnswerEntry(
                question_id=1,
                q_type="single_choice",
                stem=_rich_doc("1+1=?"),
                options=options,
                answer={"kind": "single_choice", "correct": "opt_a"},
                thinking=None,
                analysis=_rich_doc("因为 1+1=2"),
                summary=None,
            ),
        ],
    )
    return CompositionExportDoc(
        title="导出冒烟测试",
        nodes=[
            ExportHeadingNode(level=2, content=_rich_doc("第一节")),
            ExportRichTextNode(content=_rich_doc("说明文字")),
            question,
            ExportPageBreakNode(),
            module,
        ],
    )


def test_composition_docx_renderer_produces_nonempty_file():
    path = CompositionDocxRenderer().render(_sample_doc())
    try:
        assert os.path.exists(path)
        assert os.path.getsize(path) > 0
    finally:
        os.remove(path)


def test_composition_docx_renderer_does_not_print_title_in_body():
    # Composition.title 只用于导出文件名,正文必须只来自画布节点。
    path = CompositionDocxRenderer().render(_sample_doc())
    try:
        document = Document(path)
        paragraph_texts = [p.text for p in document.paragraphs]
        assert "导出冒烟测试" not in paragraph_texts
        assert paragraph_texts[0] == "第一节"
    finally:
        os.remove(path)


def test_composition_docx_renderer_applies_heading_text_align():
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = CompositionExportDoc(
        title="t",
        nodes=[ExportHeadingNode(level=2, content=_rich_doc_aligned("居中标题", "center"))],
    )
    path = CompositionDocxRenderer().render(doc)
    try:
        document = Document(path)
        assert document.paragraphs[0].text == "居中标题"
        assert document.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    finally:
        os.remove(path)


def test_composition_latex_renderer_produces_nonempty_zip_with_tex():
    path = CompositionLatexRenderer().render(_sample_doc())
    try:
        assert os.path.exists(path)
        with zipfile.ZipFile(path) as zf:
            names = zf.namelist()
            tex_names = [n for n in names if n.endswith(".tex")]
            assert len(tex_names) == 1
            tex = zf.read(tex_names[0]).decode("utf-8")
        assert "\\begin{document}" in tex
        assert "\\end{document}" in tex
        assert "\\clearpage" in tex
        assert "\\begin{multicols}{2}" in tex
        # Composition.title 只用于文件名(.tex 文件名),正文不包含标题行。
        assert "导出冒烟测试" not in tex
    finally:
        os.remove(path)


def test_composition_registry_resolves_both_formats():
    assert isinstance(composition_renderer_for(OutputFormat.DOCX), CompositionDocxRenderer)
    assert isinstance(composition_renderer_for(OutputFormat.LATEX), CompositionLatexRenderer)


def test_assembled_snapshot_renders_end_to_end_without_error():
    snapshot = {
        "schema_version": 2,
        "composition_id": 1,
        "source_revision": 1,
        "title": "端到端冒烟",
        "subject_id": 1,
        "finalized_at": "2026-01-01T00:00:00",
        "numbering_enabled": True,
        "scoring_enabled": True,
        "question_display": {"answer": True, "thinking": False, "analysis": False, "summary": False},
        "nodes": [
            {
                "id": "q1", "parent_id": None, "slot": None, "position": 0,
                "node_kind": "block", "node_type": "question", "schema_version": 1,
                "question_id": 1, "question_revision": 1,
                "props": {"number": "1", "score": 3},
                "question": {
                    "id": 1, "content_revision": 1, "content_schema_version": 2,
                    "q_type": "single_choice", "content": _rich_doc("Q1"),
                    "options": [{"id": "opt_a", "label": "A", "content": _rich_doc("A")}],
                    "answer": {"kind": "single_choice", "correct": "opt_a"},
                    "thinking": None, "analysis": None, "summary": None,
                    "difficulty": 3, "source": None,
                },
            },
        ],
    }
    doc = CompositionAssembler().assemble(snapshot)
    for fmt in (OutputFormat.DOCX, OutputFormat.LATEX):
        path = composition_renderer_for(fmt).render(doc)
        try:
            assert os.path.getsize(path) > 0
        finally:
            os.remove(path)

"""Validate a template-free "backend renders Tiptap JSON directly to docx" export pipeline.

No HTML intermediate, no --reference-doc template: walks the Tiptap JSON tree and calls
python-docx APIs directly (paragraph alignment, bold/italic/underline/sup/sub runs, lists,
images). Formulas are the one piece still borrowed from pandoc: each LaTeX fragment is
converted in isolation (markdown+tex_math_dollars -> docx) purely to extract the resulting
<m:oMath> OMML XML, which is then spliced into our own paragraph.

Usage:
    uv run python scripts/test_tiptap_export.py [path/to/tiptap.json]
"""
import copy
import json
import logging
import sys
import tempfile
import zipfile
from pathlib import Path

import pypandoc
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
logger = logging.getLogger(__name__)

M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
ALIGN_MAP = {
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}
LIST_STYLES = {"bulletList": "List Bullet", "orderedList": "List Number"}


def latex_to_omml(latex: str, display: bool) -> etree._Element:
    """Convert a single LaTeX formula to a standalone <m:oMath> element via a throwaway pandoc call."""
    markdown = f"$${latex}$$" if display else f"${latex}$"
    with tempfile.TemporaryDirectory() as tmp:
        docx_path = Path(tmp) / "formula.docx"
        pypandoc.convert_text(
            markdown, "docx", format="markdown+tex_math_dollars",
            outputfile=str(docx_path), extra_args=["--standalone"],
        )
        with zipfile.ZipFile(docx_path) as zf:
            xml_bytes = zf.read("word/document.xml")
    root = etree.fromstring(xml_bytes)
    math_el = root.find(f".//{{{M_NS}}}oMath")
    if math_el is None:
        raise ValueError(f"pandoc produced no OMML for latex: {latex!r}")
    return copy.deepcopy(math_el)


def apply_mark(run, mark: dict):
    mark_type = mark.get("type")
    if mark_type == "bold":
        run.bold = True
    elif mark_type == "italic":
        run.italic = True
    elif mark_type == "underline":
        run.underline = True
    elif mark_type == "superscript":
        run.font.superscript = True
    elif mark_type == "subscript":
        run.font.subscript = True
    else:
        logger.warning("unknown mark type %r, ignored", mark_type)


def add_inline(paragraph, node: dict):
    node_type = node.get("type")
    if node_type == "text":
        run = paragraph.add_run(node.get("text", ""))
        for mark in node.get("marks", []):
            apply_mark(run, mark)
        return
    if node_type == "hardBreak":
        paragraph.add_run().add_break()
        return
    if node_type == "inlineMath":
        latex = node.get("attrs", {}).get("latex", "")
        try:
            paragraph._p.append(latex_to_omml(latex, display=False))
        except Exception:
            logger.warning("failed to convert inline latex %r, falling back to plain text", latex)
            paragraph.add_run(latex)
        return
    logger.warning("unknown inline node type %r, skipped", node_type)


def add_block(document: Document, node: dict, list_style: str | None = None):
    node_type = node.get("type")
    if node_type == "paragraph":
        p = document.add_paragraph(style=list_style)
        align = node.get("attrs", {}).get("textAlign")
        if align in ALIGN_MAP:
            p.alignment = ALIGN_MAP[align]
        for child in node.get("content", []):
            add_inline(p, child)
        return
    if node_type == "blockMath":
        latex = node.get("attrs", {}).get("latex", "")
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            p._p.append(latex_to_omml(latex, display=True))
        except Exception:
            logger.warning("failed to convert block latex %r, falling back to plain text", latex)
            p.add_run(latex)
        return
    if node_type in LIST_STYLES:
        style = LIST_STYLES[node_type]
        for item in node.get("content", []):
            for child in item.get("content", []):
                add_block(document, child, list_style=style)
        return
    if node_type == "image":
        src = node.get("attrs", {}).get("src", "")
        try:
            document.add_picture(src)
        except Exception:
            logger.warning("failed to embed image %r, skipped", src)
        return

    # Unknown node types are dropped instead of blindly passed through.
    logger.warning("unknown block node type %r, skipped", node_type)


def render_document(doc: dict) -> Document:
    document = Document()
    for node in doc.get("content", []):
        add_block(document, node)
    return document


def inspect_docx(docx_path: Path):
    with zipfile.ZipFile(docx_path) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
    return {
        "omath_count": xml.count("<m:oMath ") + xml.count("<m:oMath>"),
        "jc_count": xml.count("<w:jc "),
    }


def main():
    json_path = Path(sys.argv[1]) if len(sys.argv) > 1 else Path(__file__).resolve().parent.parent.parent / "test.json"
    out_dir = Path(__file__).resolve().parent.parent / "static" / "media" / "export-test"
    out_dir.mkdir(parents=True, exist_ok=True)

    doc = json.loads(json_path.read_text(encoding="utf-8"))
    logger.info("loaded %s", json_path)

    document = render_document(doc)
    docx_path = out_dir / "test_export_native.docx"
    document.save(str(docx_path))

    stats = inspect_docx(docx_path)
    logger.info(
        "wrote docx: %s (%d bytes), oMath elements: %d, paragraph alignments set: %d",
        docx_path, docx_path.stat().st_size, stats["omath_count"], stats["jc_count"],
    )


if __name__ == "__main__":
    main()

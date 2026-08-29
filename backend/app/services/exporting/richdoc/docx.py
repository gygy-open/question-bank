"""RichDoc → python-docx(直接构建,不经 Markdown / reference-doc)。

- 行内:text+marks(bold/italic/underline/sup/sub)、inlineMath(OMML)、hardBreak、image、blank。
- 块级:paragraph(textAlign)、blockMath、bullet/ordered list、image、table。
- 公式走 richdoc.omml(in-process,无 pandoc);单条失败退化为纯文本 run。
- 未知/未来节点绝不丢字符:优先递归 content,否则退化为纯文本 run。
"""

from __future__ import annotations

import logging
import math
from typing import Any, Optional

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Inches

from app.services.exporting.images import ImageResolver
from app.services.exporting.richdoc.omml import latex_to_omml

logger = logging.getLogger(__name__)

_ALIGN = {
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "left": WD_ALIGN_PARAGRAPH.LEFT,
}
_LIST_STYLE = {"bulletList": "List Bullet", "orderedList": "List Number"}

_BLANK_WIDTH_MIN_EM = 2
_BLANK_WIDTH_MAX_EM = 30
_BLANK_WIDTH_DEFAULT_EM = 4
_EMU_PER_PX = 9525  # 914400 EMU/in ÷ 96 px/in


class DocxRichRenderer:
    """把 RichDoc 渲染进给定的 python-docx 容器(Document 或表格单元格)。"""

    def __init__(self, images: Optional[ImageResolver] = None):
        self.images = images or ImageResolver()

    # --- 块级 ---------------------------------------------------------- #
    def render_doc(self, parent: Any, doc: Any) -> None:
        if not isinstance(doc, dict):
            return
        self.add_blocks(parent, doc.get("content") or [])

    def add_blocks(self, parent: Any, nodes: list[Any], list_style: Optional[str] = None) -> None:
        for node in nodes:
            if isinstance(node, dict):
                self.add_block(parent, node, list_style)

    def add_block(self, parent: Any, node: dict[str, Any], list_style: Optional[str] = None) -> None:
        t = node.get("type")

        if t == "paragraph":
            p = parent.add_paragraph(style=list_style) if list_style else parent.add_paragraph()
            align = (node.get("attrs") or {}).get("textAlign")
            if align in _ALIGN:
                p.alignment = _ALIGN[align]
            for child in node.get("content") or []:
                if isinstance(child, dict):
                    self.add_inline(p, child)
            return

        if t == "blockMath":
            latex = (node.get("attrs") or {}).get("latex", "")
            p = parent.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._add_math(p, latex, display=True)
            return

        if t in _LIST_STYLE:
            style = _LIST_STYLE[t]
            for item in node.get("content") or []:
                if isinstance(item, dict):
                    self.add_blocks(parent, item.get("content") or [], list_style=style)
            return

        if t == "image":
            p = parent.add_paragraph()
            self._add_image(p, node)
            return

        if t == "table":
            self._add_table(parent, node)
            return

        if t == "blockquote":
            self._add_blockquote(parent, node.get("content") or [])
            return

        if t == "codeBlock":
            self._add_code_block(parent, node)
            return

        if t == "horizontalRule":
            self._add_horizontal_rule(parent)
            return

        # 未知/未来块:优先递归,否则退化为纯文本段落。
        if node.get("content"):
            self.add_blocks(parent, node.get("content") or [], list_style)
        elif node.get("text"):
            parent.add_paragraph().add_run(str(node.get("text")))

    def _collect_text(self, node: dict[str, Any]) -> str:
        """收集块内文本（代码块用），hardBreak 视为换行。"""
        parts: list[str] = []
        for child in node.get("content") or []:
            if not isinstance(child, dict):
                continue
            if child.get("type") == "text":
                parts.append(str(child.get("text") or ""))
            elif child.get("type") == "hardBreak":
                parts.append("\n")
            else:
                parts.append(self._collect_text(child))
        return "".join(parts)

    def _add_blockquote(self, parent: Any, nodes: list[Any]) -> None:
        for node in nodes:
            if not isinstance(node, dict):
                continue
            if node.get("type") == "paragraph":
                try:
                    p = parent.add_paragraph(style="Quote")
                except (KeyError, ValueError):
                    p = parent.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.5)
                for child in node.get("content") or []:
                    if isinstance(child, dict):
                        self.add_inline(p, child)
            else:
                self.add_block(parent, node)

    def _add_code_block(self, parent: Any, node: dict[str, Any]) -> None:
        p = parent.add_paragraph()
        for i, line in enumerate(self._collect_text(node).split("\n")):
            if i:
                p.add_run().add_break()
            run = p.add_run(line)
            run.font.name = "Consolas"

    def _add_horizontal_rule(self, parent: Any) -> None:
        # 底边框空段落模拟分隔线（python-docx 无原生 HR）。
        p = parent.add_paragraph()
        p_pr = p._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")
        borders.append(bottom)
        p_pr.append(borders)

    # --- 行内 ---------------------------------------------------------- #
    def add_inline(self, paragraph: Any, node: dict[str, Any]) -> None:
        t = node.get("type")

        if t == "text":
            run = paragraph.add_run(node.get("text", ""))
            self._apply_marks(run, node.get("marks") or [])
            return

        if t == "hardBreak":
            paragraph.add_run().add_break()
            return

        if t == "inlineMath":
            self._add_math(paragraph, (node.get("attrs") or {}).get("latex", ""), display=False)
            return

        if t == "image":
            self._add_image(paragraph, node)
            return

        if t == "blank":
            run = paragraph.add_run("\u00a0" * self._blank_len(node))
            run.underline = True
            return

        # 未知/未来行内:递归 content 或退化为文本 run。
        if node.get("content"):
            for child in node.get("content") or []:
                if isinstance(child, dict):
                    self.add_inline(paragraph, child)
        elif node.get("text"):
            paragraph.add_run(str(node.get("text")))

    # --- 辅助 ---------------------------------------------------------- #
    def _apply_marks(self, run: Any, marks: list[Any]) -> None:
        for mark in marks:
            if not isinstance(mark, dict):
                continue
            mt = mark.get("type")
            if mt == "bold":
                run.bold = True
            elif mt == "italic":
                run.italic = True
            elif mt == "underline":
                run.underline = True
            elif mt == "superscript":
                run.font.superscript = True
            elif mt == "subscript":
                run.font.subscript = True

    def _add_math(self, paragraph: Any, latex: str, display: bool) -> None:
        try:
            paragraph._p.append(latex_to_omml(latex, display=display))
        except Exception:
            logger.warning("latex→OMML failed for %r, falling back to text", latex)
            paragraph.add_run(latex)

    def _add_image(self, paragraph: Any, node: dict[str, Any]) -> None:
        attrs = node.get("attrs") or {}
        src = str(attrs.get("src", ""))
        path = self.images.resolve(src)
        if path is None:
            alt = str(attrs.get("alt", "") or "")
            if alt:
                paragraph.add_run(alt)
            return
        width = attrs.get("width")
        kwargs = {}
        if isinstance(width, (int, float)) and not isinstance(width, bool) and math.isfinite(float(width)) and width > 0:
            kwargs["width"] = Emu(int(float(width) * _EMU_PER_PX))
        try:
            paragraph.add_run().add_picture(str(path), **kwargs)
        except Exception:
            logger.warning("failed to embed image %s, skipped", src)

    def _add_table(self, parent: Any, node: dict[str, Any]) -> None:
        rows = [r for r in (node.get("content") or []) if isinstance(r, dict) and r.get("type") == "tableRow"]
        if not rows:
            return
        grid = []
        for row in rows:
            grid.append([c for c in (row.get("content") or []) if isinstance(c, dict)])
        ncols = max(len(r) for r in grid)
        table = parent.add_table(rows=len(grid), cols=ncols)
        table.style = "Table Grid"
        for i, cells in enumerate(grid):
            for j in range(ncols):
                if j < len(cells):
                    self.add_blocks(table.cell(i, j), cells[j].get("content") or [])

    def _blank_len(self, node: dict[str, Any]) -> int:
        value = (node.get("attrs") or {}).get("widthEm")
        if isinstance(value, bool) or not isinstance(value, (int, float)) or not math.isfinite(float(value)):
            width = _BLANK_WIDTH_DEFAULT_EM
        else:
            width = min(_BLANK_WIDTH_MAX_EM, max(_BLANK_WIDTH_MIN_EM, float(value)))
        return max(2, round(width))

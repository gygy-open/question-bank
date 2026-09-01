"""CompositionExportDoc → python-docx 直接构建(结构与 Paper 的 DocxRenderer 同构)。"""

from __future__ import annotations

import os
import tempfile
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from app.services.exporting.answer import answer_spec_to_inline
from app.services.exporting.composition_contracts import (
    CompositionExportDoc,
    CompositionExportNode,
    ExportAnswerEntry,
    ExportAnswerSpaceNode,
    ExportHeadingNode,
    ExportOption,
    ExportPageBreakNode,
    ExportQuestionDetailsNode,
    ExportQuestionNode,
    ExportRichTextNode,
    RichDoc,
)
from app.services.exporting.images import ImageResolver
from app.services.exporting.richdoc.docx import DocxRichRenderer

_HEADING_SIZES = {1: 18, 2: 16, 3: 14, 4: 12}
_LABELS = (("thinking", "【思路】"), ("analysis", "【解析】"), ("summary", "【总结】"))# 与前端画布一致：有题号时题干悬挂缩进（折行对齐首行文字起点），选项跟随同一缩进；分值只挤在首行,不参与缩进量。
_NUMBER_INDENT_PT = 21# 与 richdoc/docx.py 的 _ALIGN 保持一致（那边是模块私有常量，不跨模块导入）。
_ALIGN = {
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "left": WD_ALIGN_PARAGRAPH.LEFT,
}


def _format_score(score: float) -> str:
    return f"{score:g}"


def _set_bottom_border(paragraph: Any) -> None:
    """给段落加一条下边框，作为作答横线。"""
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    borders.append(bottom)
    p_pr.append(borders)


class CompositionDocxRenderer:
    ext = "docx"

    def __init__(self) -> None:
        self.rich = DocxRichRenderer(ImageResolver())

    def render(self, doc: CompositionExportDoc) -> str:
        # doc.title 只用于导出文件名（见 composition_registry/API 层），正文完全按画布节点渲染,
        # 不额外插入标题行。
        document = Document()
        self._set_cjk_font(document)

        for node in doc.nodes:
            self._add_node(document, node)

        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        document.save(path)
        return path

    # --- 节点 dispatch --------------------------------------------------- #
    def _add_node(self, document: Document, node: CompositionExportNode) -> None:
        if isinstance(node, ExportRichTextNode):
            self.rich.render_doc(document, node.content)
        elif isinstance(node, ExportHeadingNode):
            self._add_heading(document, node)
        elif isinstance(node, ExportQuestionNode):
            self._add_question(document, node)
        elif isinstance(node, ExportPageBreakNode):
            document.add_page_break()
        elif isinstance(node, ExportAnswerSpaceNode):
            self._add_answer_space(document, node)
        elif isinstance(node, ExportQuestionDetailsNode):
            self._add_question_details(document, node)

    def _add_answer_space(self, document: Document, node: ExportAnswerSpaceNode) -> None:
        # 逐行加空段落;lined 样式给每段落加下边框形成答题横线。
        for _ in range(max(1, node.lines)):
            p = document.add_paragraph()
            if node.style == "lined":
                _set_bottom_border(p)

    def _add_heading(self, document: Document, node: ExportHeadingNode) -> None:
        # heading content 恒为单段落纯行内内容(schema 已保证),直接抽取该段落渲染并整体加粗放大，
        # 同时保留编辑器中设置的段落对齐(textAlign)。
        p = document.add_paragraph()
        blocks = (node.content or {}).get("content") or []
        first = blocks[0] if blocks else None
        if isinstance(first, dict):
            align = (first.get("attrs") or {}).get("textAlign")
            if align in _ALIGN:
                p.alignment = _ALIGN[align]
            for child in first.get("content") or []:
                if isinstance(child, dict):
                    self.rich.add_inline(p, child)
        size = _HEADING_SIZES.get(node.level, 12)
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(size)

    def _add_question(self, document: Document, q: ExportQuestionNode) -> None:
        # 题号(粗体)+ 分值(斜体)同挤在题干首段前,而非各占一行,与编辑器内联展示口径一致。
        p = document.add_paragraph()
        if q.number:
            p.paragraph_format.left_indent = Pt(_NUMBER_INDENT_PT)
            p.paragraph_format.first_line_indent = Pt(-_NUMBER_INDENT_PT)
            p.add_run(f"{q.number}. ").bold = True
        if q.score is not None:
            p.add_run(f"（{_format_score(q.score)} 分）").italic = True
        self._append_richdoc(p, document, q.stem)
        self._add_options(document, q.options, q.option_columns, indent=bool(q.number))
        self._add_inline_fields(document, q.answer, q.thinking, q.analysis, q.summary, q.options)

    def _add_options(
        self, document: Document, options: list[ExportOption], columns: int, indent: bool = False,
    ) -> None:
        if not options:
            return
        columns = max(1, columns)
        rows = -(-len(options) // columns)  # ceil div
        table = document.add_table(rows=rows, cols=columns)
        if indent:
            self._set_table_indent(table, _NUMBER_INDENT_PT)
        for idx, opt in enumerate(options):
            r, c = divmod(idx, columns)
            self._add_prefixed(table.cell(r, c), f"{opt.label}. ", opt.content, bold_prefix=False)

    def _set_table_indent(self, table: Any, pt: float) -> None:
        indent = OxmlElement("w:tblInd")
        indent.set(qn("w:w"), str(int(Pt(pt).twips)))
        indent.set(qn("w:type"), "dxa")
        table._tbl.tblPr.append(indent)

    def _add_question_details(self, document: Document, node: ExportQuestionDetailsNode) -> None:
        label = "答案汇总（此前题目）" if node.scope == "before" else "答案汇总（全篇）"
        p = document.add_paragraph()
        p.add_run(label).bold = True
        for child in node.children:
            if isinstance(child, ExportHeadingNode):
                self._add_heading(document, child)
            elif isinstance(child, ExportRichTextNode):
                self.rich.render_doc(document, child.content)
            elif isinstance(child, ExportAnswerEntry):
                self._add_answer_entry(document, child)

    def _add_answer_entry(self, document: Document, entry: ExportAnswerEntry) -> None:
        # 题号(如有)与首个可见字段同段落起笔,并对每个字段段落施加悬挂缩进,使折行对齐首行文字
        # 起点而非题号下方/左边距——即 word 编号列表的观感。
        fields: list[tuple[str, Any]] = []
        if entry.answer is not None:
            fields.append(("answer", entry.answer))
        for value, label in zip((entry.thinking, entry.analysis, entry.summary), (l for _, l in _LABELS)):
            if value:
                fields.append((label, value))

        for idx, (label, value) in enumerate(fields):
            p = document.add_paragraph()
            if idx == 0 and entry.number:
                p.add_run(f"{entry.number}. ").bold = True
            if label == "answer":
                p.add_run("【答案】").bold = True
                for inline_node in answer_spec_to_inline(value, entry.options):
                    self.rich.add_inline(p, inline_node)
            else:
                p.add_run(label).bold = True
                self._append_richdoc(p, document, value)
            p.paragraph_format.left_indent = Pt(21)
            p.paragraph_format.first_line_indent = Pt(-21)

    # --- 内联答案/思路/解析/小结 ------------------------------------------ #
    def _add_inline_fields(
        self,
        document: Document,
        answer: Any,
        thinking: RichDoc,
        analysis: RichDoc,
        summary: RichDoc,
        options: list[ExportOption],
    ) -> None:
        if answer is not None:
            p = document.add_paragraph()
            p.add_run("【答案】").bold = True
            for inline_node in answer_spec_to_inline(answer, options):
                self.rich.add_inline(p, inline_node)
        for value, label in zip((thinking, analysis, summary), (l for _, l in _LABELS)):
            if value:
                self._add_prefixed(document, label, value, bold_prefix=True)

    # --- 通用:前缀 + RichDoc(首段内联紧跟前缀,其余块正常渲染) -------------- #
    def _add_prefixed(self, container: Any, prefix: str, richdoc: RichDoc, bold_prefix: bool = True) -> None:
        default_p = container.paragraphs[0] if getattr(container, "paragraphs", None) else None
        p = container.add_paragraph()
        if prefix:
            p.add_run(prefix).bold = bold_prefix
        self._append_richdoc(p, container, richdoc)
        # 表格单元格自带一个空默认段落;写入内容后清掉它,避免每格顶部多一行空白。
        if default_p is not None and not default_p.runs and len(container.paragraphs) > 1:
            default_p._element.getparent().remove(default_p._element)

    def _append_richdoc(self, p: Any, container: Any, richdoc: RichDoc) -> None:
        # 首段内联内容紧跟在 p 已有的前缀 run(s) 之后,其余块沿用 container 正常追加。
        blocks = (richdoc or {}).get("content") or [] if isinstance(richdoc, dict) else []
        first = blocks[0] if blocks else None
        if isinstance(first, dict) and first.get("type") == "paragraph":
            for child in first.get("content") or []:
                if isinstance(child, dict):
                    self.rich.add_inline(p, child)
            self.rich.add_blocks(container, blocks[1:])
        else:
            self.rich.add_blocks(container, blocks)

    def _set_cjk_font(self, document: Document) -> None:
        style = document.styles["Normal"]
        style.font.size = Pt(12)
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), "SimSun")

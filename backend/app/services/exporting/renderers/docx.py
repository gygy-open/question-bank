"""DOCX 渲染器:ExportDoc → python-docx 直接构建(无 Markdown、无 pandoc)。"""

from __future__ import annotations

import logging
import os
import tempfile
from typing import Any, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

from app.services.exporting.answer import answer_spec_to_inline
from app.services.exporting.contracts import ExportDoc, ExportOption, ExportQuestion, ExportSection
from app.services.exporting.images import ImageResolver
from app.services.exporting.richdoc.docx import DocxRichRenderer

logger = logging.getLogger(__name__)

_LABELS = [
    ("thinking", "【分析】"),
    ("analysis", "【解析】"),
    ("summary", "【总结】"),
]


class DocxRenderer:
    ext = "docx"

    def __init__(self) -> None:
        self.rich = DocxRichRenderer(ImageResolver())

    def render(self, doc: ExportDoc) -> str:
        document = Document()
        self._set_cjk_font(document)

        heading = document.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run(doc.title)
        run.bold = True
        run.font.size = Pt(18)

        for section in doc.sections:
            self._add_section(document, section, body=True)

        if doc.has_appendix:
            document.add_page_break()
            title = document.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = title.add_run("参考答案与解析")
            r.bold = True
            r.font.size = Pt(15)
            for section in doc.appendix:
                self._add_section(document, section, body=False)

        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        document.save(path)
        return path

    def _add_section(self, document: Document, section: ExportSection, body: bool) -> None:
        if section.title:
            p = document.add_paragraph()
            r = p.add_run(section.title)
            r.bold = True
            r.font.size = Pt(14)
        for q in section.questions:
            if body:
                self._add_question_body(document, q)
            else:
                self._add_question_details(document, q, appendix=True)

    def _add_question_body(self, document: Document, q: ExportQuestion) -> None:
        self._add_prefixed(document, f"{q.number}. ", q.stem)
        for opt in q.options:
            self._add_prefixed(document, f"{opt.label}. ", opt.content)
        if q.reserve_space:
            for _ in range(4):
                document.add_paragraph()
        self._add_question_details(document, q, appendix=False)

    def _add_question_details(self, document: Document, q: ExportQuestion, appendix: bool) -> None:
        if appendix:
            p = document.add_paragraph()
            p.add_run(f"{q.number}. ").bold = True
        if q.answer is not None:
            p = document.add_paragraph()
            p.add_run("【答案】").bold = True
            for node in answer_spec_to_inline(q.answer, q.options):
                self.rich.add_inline(p, node)
        for attr, label in _LABELS:
            value = getattr(q, attr)
            if value:
                self._add_prefixed(document, label, value, bold_prefix=True)
        if q.source:
            p = document.add_paragraph()
            p.add_run("【来源】").bold = True
            p.add_run(q.source)

    def _add_prefixed(
        self,
        document: Document,
        prefix: str,
        doc: Any,
        bold_prefix: bool = True,
    ) -> None:
        """渲染一个 RichDoc,并把 prefix 注入首个段落;首块非段落时另起前缀段。"""
        blocks = (doc or {}).get("content") or [] if isinstance(doc, dict) else []
        first = blocks[0] if blocks else None
        if isinstance(first, dict) and first.get("type") == "paragraph":
            p = document.add_paragraph()
            r = p.add_run(prefix)
            r.bold = bold_prefix
            for child in first.get("content") or []:
                if isinstance(child, dict):
                    self.rich.add_inline(p, child)
            self.rich.add_blocks(document, blocks[1:])
        else:
            p = document.add_paragraph()
            r = p.add_run(prefix)
            r.bold = bold_prefix
            self.rich.add_blocks(document, blocks)

    def _set_cjk_font(self, document: Document) -> None:
        style = document.styles["Normal"]
        style.font.size = Pt(12)
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), "SimSun")
